#!/usr/bin/env python3
"""Export PDF to txt/png/docx/xlsx using PyMuPDF and office libraries."""

import argparse
import os
import sys

import fitz


def export_txt(doc, out):
    with open(out, "w", encoding="utf-8") as f:
        for page in doc:
            f.write(page.get_text())
            f.write("\n\n")
    print(f"Exported text -> {out}")


def export_png(doc, outdir):
    os.makedirs(outdir, exist_ok=True)
    for i, page in enumerate(doc, 1):
        pix = page.get_pixmap(dpi=150)
        path = os.path.join(outdir, f"page-{i}.png")
        pix.save(path)
    print(f"Exported {len(doc)} page image(s) -> {outdir}")


def export_docx(doc, out):
    from docx import Document

    d = Document()
    for i, page in enumerate(doc, 1):
        d.add_heading(f"Page {i}", level=1)
        text = page.get_text().strip()
        if text:
            d.add_paragraph(text)
    d.save(out)
    print(f"Exported Word -> {out}")


def export_xlsx(doc, out):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "PDF Export"
    row = 1
    for i, page in enumerate(doc, 1):
        ws.cell(row=row, column=1, value=f"Page {i}")
        row += 1
        for line in page.get_text().splitlines():
            if line.strip():
                ws.cell(row=row, column=1, value=line)
                row += 1
        row += 1
    wb.save(out)
    print(f"Exported Excel -> {out}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--format", required=True, choices=["txt", "png", "docx", "xlsx"])
    parser.add_argument("--output", "-o", required=True)
    parser.add_argument("file")
    args = parser.parse_args()

    doc = fitz.open(args.file)
    if args.format == "txt":
        export_txt(doc, args.output)
    elif args.format == "png":
        export_png(doc, args.output)
    elif args.format == "docx":
        export_docx(doc, args.output)
    elif args.format == "xlsx":
        export_xlsx(doc, args.output)
    doc.close()
    return 0


if __name__ == "__main__":
    sys.exit(main())
